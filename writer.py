from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

MONTHS = {
    1: 'января', 2: 'февраля', 3: 'марта', 4: 'апреля',
    5: 'мая', 6: 'июня', 7: 'июля', 8: 'августа',
    9: 'сентября', 10: 'октября', 11: 'ноября', 12: 'декабря'
}

def format_russian_date(date_str):
    if not date_str:
        return "«__» _________ 20__ г."

    try:
        dt = datetime.datetime.strptime(date_str, "%Y-%m-%d")
        month_name = MONTHS[dt.month]
        return f"«{dt.day}» {month_name} {dt.year} г."
    except ValueError:
        return date_str

def set_font(run, size=12, bold=False, underline=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.underline = underline

def remove_borders(table):
    tbl = table._tbl
    for cell in tbl.iter_tcs():
        tcPr = cell.tcPr
        tcBorders = OxmlElement('w:tcBorders')
        for border in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            node = OxmlElement(f'w:{border}')
            node.set(qn('w:val'), 'nil')
            tcBorders.append(node)
        tcPr.append(tcBorders)

def save(data, filepath):
    doc = Document()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("СОГЛАСИЕ\nна обработку персональных данных")
    set_font(run, size=14, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)

    set_font(p.add_run("Я, "))
    set_font(p.add_run(f" {data['fio']} "), underline=True)
    p.add_run("\n")

    set_font(p.add_run("проживающий(ая) по адресу: "))
    set_font(p.add_run(f" {data['address']} "), underline=True)
    p.add_run("\n")

    p.add_run("_" * 80)

    p = doc.add_paragraph()
    set_font(p.add_run("документ, удостоверяющий личность: серия "))
    set_font(p.add_run(f" {data['series']} "), underline=True)
    set_font(p.add_run(" № "))
    set_font(p.add_run(f" {data['number']} "), underline=True)
    set_font(p.add_run(",\nвыдан "))
    set_font(p.add_run(f" {data['issued']} "), underline=True)

    body_text = (
        "в соответствии с требованиями статьи 9 Федерального закона от 27 июля 2006 "
        "г. № 152-ФЗ «О персональных данных», свободно, своей волей и в своем "
        "интересе даю согласие на обработку (любое действие (операцию) или совокупность "
        "действий (операций), совершаемых с использованием средств автоматизации "
        "или без использования таких средств с персональными данными, включая "
        "сбор, запись, систематизацию, накопление, хранение, уточнение (обновление, "
        "изменение), извлечение, использование, передачу (распространение, "
        "предоставление, доступ), обезличивание, блокирование, удаление, "
        "уничтожение) следующих персональных данных:\n"
        "   - фамилия, имя, отчество;\n"
        "   - паспортные данные (серия, номер, кем и когда выдан);\n"
        "   - адрес регистрации и фактического проживания;\n"
        "   - номер контактного телефона/email.\n\n"
        "Вышеуказанные персональные данные предоставляю для обработки в целях "
        "осуществления возврата ошибочно уплаченных (взысканных) платежей.\n"
        "Настоящее согласие действует на период до истечения сроков хранения "
        "соответствующей информации или документов."
    )

    p = doc.add_paragraph(body_text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    for run in p.runs:
        set_font(run)

    p = doc.add_paragraph("Отзыв согласия осуществляется в соответствии с законодательством Российской Федерации.")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(24)
    set_font(p.runs[0])

    table = doc.add_table(rows=1, cols=2)
    remove_borders(table)

    table.columns[0].width = Cm(11)
    table.columns[1].width = Cm(6)

    row1 = table.rows[0]

    p1 = row1.cells[0].paragraphs[0]
    set_font(p1.add_run(f"{data['fio']}"), underline=True)

    p2 = row1.cells[1].paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    formatted_date = format_russian_date(data['date'])
    set_font(p2.add_run(formatted_date))

    row2 = table.add_row()
    p3 = row2.cells[0].paragraphs[0]
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p3.add_run("(фамилия, имя, отчество)"), size=8)

    row3 = table.add_row()
    p4 = row3.cells[0].paragraphs[0]
    p4.paragraph_format.space_before = Pt(12)
    set_font(p4.add_run("_________________________"), underline=False)

    row4 = table.add_row()
    p5 = row4.cells[0].paragraphs[0]
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p5.add_run("(подпись)"), size=8)

    doc.save(filepath)
